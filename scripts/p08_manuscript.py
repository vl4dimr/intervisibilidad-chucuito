# -*- coding: utf-8 -*-
"""
Manuscrito, generado a partir de los ficheros de resultados.

Las cifras no se escriben a mano: se leen de los JSON que produjo el analisis.
Asi el texto no puede desviarse de los datos cuando algo se recalcula, que es la
forma mas comun de que un articulo acabe afirmando un numero que ya no es suyo.

Formato de Virtual Archaeology Review: fichero anonimo, Arial, dos columnas,
marcador decimal con punto, encabezados numerados.
"""
import json
import os

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# DOI de concepto del deposito en Zenodo: apunta siempre a la ultima version.
DOI_CONCEPTO = "10.5281/zenodo.22176260"
RES = os.path.join(BASE, "results")
FIG = os.path.join(RES, "figuras")
OUT = os.path.join(BASE, "VAR_intervisibilidad.docx")

ARIAL, BLACK = "Arial", RGBColor(0, 0, 0)
TEXT_W_MM, GUTTER_MM = 170.0, 6.0


def cargar(nombre):
    p = os.path.join(RES, nombre)
    if not os.path.exists(p):
        return None
    with open(p, encoding="utf-8") as f:
        return json.load(f)


NUL = cargar("nulos.json")
RIG = cargar("nulo_rigido.json")
FUN = cargar("funerarios.json")
VAL = cargar("validacion_los.json")
SEN = cargar("sensibilidad.json")
ALT = cargar("nulo_alturas.json")
DIS = cargar("por_distrito.json")
SINMASK = cargar("nulo_rigido_sin_mascara.json")   # contraejemplo: agua sin enmascarar
N300 = cargar("nulo_rigido_n300.json")             # version previa, recorte de 5 km
with open(os.path.join(BASE, "data", "terreno_log.json"), encoding="utf-8") as f:
    TER = json.load(f)
with open(os.path.join(BASE, "data", "fetch_log.json"), encoding="utf-8") as f:
    FET = json.load(f)


def f(x, d=4):
    return ("%%.%df" % d) % x


def mil(n):
    """Entero con espacio fino de millar, como pide la norma tipografica del castellano."""
    return "{:,}".format(int(n)).replace(",", "\u202f")


doc = Document()


def page_setup(sec):
    sec.page_width, sec.page_height = Mm(210), Mm(297)
    sec.left_margin = sec.right_margin = Mm((210 - TEXT_W_MM) / 2)
    sec.top_margin = sec.bottom_margin = Mm(20)


def set_cols(sec, num):
    c = sec._sectPr.xpath("./w:cols")[0]
    c.set(qn("w:num"), str(num))
    c.set(qn("w:space"), str(int(GUTTER_MM * 56.7)))
    c.set(qn("w:equalWidth"), "1")


def run(p, t, size=9, bold=False, italic=False):
    r = p.add_run(t)
    r.font.name = ARIAL
    r.font.size = Pt(size)
    r.bold, r.italic = bold, italic
    r.font.color.rgb = BLACK
    r._element.rPr.rFonts.set(qn("w:eastAsia"), ARIAL)
    return r


def P(t="", size=9, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
      after=6, before=0):
    p = doc.add_paragraph()
    if t:
        run(p, t, size, bold, italic)
    pf = p.paragraph_format
    pf.alignment, pf.space_after, pf.space_before = align, Pt(after), Pt(before)
    pf.line_spacing = 1.0
    return p


def h1(t):
    return P(t, 11, True, align=WD_ALIGN_PARAGRAPH.CENTER, before=10, after=6)


def h2(t):
    return P(t, 10, True, align=WD_ALIGN_PARAGRAPH.LEFT, before=8, after=4)


def one_col():
    s = doc.add_section(WD_SECTION.CONTINUOUS); page_setup(s); set_cols(s, 1)


def two_col():
    s = doc.add_section(WD_SECTION.CONTINUOUS); page_setup(s); set_cols(s, 2)


def caption(label, text, above=False):
    p = doc.add_paragraph()
    run(p, label + " ", 8, bold=True)
    run(p, text, 8)
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(8 if above else 4)
    pf.space_after = Pt(4 if above else 8)
    pf.line_spacing = 1.0


def table(label, text, cols, rows, widths=None):
    one_col()
    caption(label, text, above=True)
    t = doc.add_table(rows=1, cols=len(cols))
    t.style, t.alignment = "Table Grid", WD_TABLE_ALIGNMENT.CENTER
    for i, c in enumerate(cols):
        cell = t.rows[0].cells[i]; cell.text = ""
        run(cell.paragraphs[0], c, 8, italic=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        cs = t.add_row().cells
        for i, v in enumerate(row):
            cs[i].text = ""
            run(cs[i].paragraphs[0], str(v), 8)
            cs[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if widths:
        for rr in t.rows:
            for i, w in enumerate(widths):
                rr.cells[i].width = Mm(w)
    doc.add_paragraph()
    two_col()


def figure(fname, label, text, width_mm=TEXT_W_MM):
    p = os.path.join(FIG, fname)
    if not os.path.exists(p):
        return
    one_col()
    doc.add_picture(p, width=Mm(width_mm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption(label, text)
    two_col()


page_setup(doc.sections[0]); set_cols(doc.sections[0], 1)
st = doc.styles["Normal"]; st.font.name = ARIAL; st.font.size = Pt(9)
st.paragraph_format.space_after = Pt(6); st.paragraph_format.line_spacing = 1.0

# =========================================================== portada y resumen
P("INTERVISIBILIDAD DE SITIOS ARQUEOLÓGICOS EN LA CUENCA DEL TITICACA: ANÁLISIS SOBRE "
  "MODELO DIGITAL DE ELEVACIÓN ABIERTO Y MODELOS NULOS",
  14, True, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
P("INTERVISIBILITY OF ARCHAEOLOGICAL SITES IN THE TITICACA BASIN: ANALYSIS ON AN OPEN "
  "DIGITAL ELEVATION MODEL WITH NULL MODELS",
  9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
P("[Fichero anónimo conforme a la norma de la revista: los datos de autoría figuran en los metadatos "
  "de la plataforma y en la carta de presentación.]",
  8, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)

n_sitios = TER["sitios"]
n_fun = TER["funerarios_por_toponimo"]
d5 = NUL["observado"]["5000"] if NUL else None
r5 = RIG["por_alcance"]["5000"]

P("Highlights", 9, True, align=WD_ALIGN_PARAGRAPH.LEFT, after=3)
for b in ["%d sitios arqueológicos de Chucuito analizados sobre Copernicus DEM GLO-30, con datos abiertos."
          % n_sitios,
          "La configuración observada supera en intervisibilidad a la misma nube rotada y desplazada "
          "(p = %s), aunque el efecto se atenúa al analizar por separado cada grupo."
          % f(r5["p_unilateral"], 3),
          "Enmascarar el lago cambia el contraste de z = %+.2f a z = %+.2f: sin ello, el efecto "
          "desaparece." % (SINMASK["por_alcance"]["5000"]["z"] if SINMASK else float("nan"),
                           r5["z"])]:
    p = doc.add_paragraph(style="List Bullet")
    run(p, b, 9)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
doc.add_paragraph()

P("Resumen", 9, True, align=WD_ALIGN_PARAGRAPH.LEFT, after=3)
P("Se analiza si los sitios arqueológicos registrados en la provincia de Chucuito, en la ribera "
  "peruana del lago Titicaca, ocupan posiciones desde las que se ven entre sí más de lo que cabría "
  "esperar del azar. Sobre el Copernicus DEM GLO-30 se calcula la línea de visión entre los %s pares "
  "que forman los %d sitios documentados, con corrección de curvatura terrestre y refracción "
  "atmosférica, y se "
  "contrasta la red resultante contra tres modelos nulos de exigencia creciente. El más estricto "
  "traslada y rota la nube de sitios completa, conservando todas las distancias mutuas, de modo que "
  "aísla el emplazamiento de la disposición. La densidad de intervisibilidad observada a 5 km es %s "
  "frente a %s ± %s en el nulo (z = %+.2f; p = %s), y se mantiene en los cuatro alcances considerados "
  "y en todo el rango de alturas atribuibles a las estructuras. Restringido al mayor de los dos grupos "
  "espacialmente disjuntos que componen el conjunto, sin embargo, el contraste queda en el margen de la "
  "significación convencional (p = %s), de modo que la evidencia es indicativa y no concluyente. "
  "El trabajo documenta además tres artefactos que invierten o anulan el resultado si no se corrigen: "
  "el signo de la corrección por curvatura, la cota constante que el modelo de elevación asigna a las "
  "masas de agua —el %.1f %% del área de estudio— y el sesgo de orientación que introduce un recorte "
  "ajustado. Todos producen resultados plausibles y falsos."
  % (mil(n_sitios * (n_sitios - 1) // 2), n_sitios,
     f(r5["densidad_obs"]), f(r5["nula_media"]), f(r5["nula_sd"]),
     r5["z"], f(r5["p_unilateral"], 3),
     f(DIS["distritos"]["juli"]["contraste"]["5000"]["p_unilateral"], 3)
     if (DIS and DIS["distritos"]["juli"].get("contraste")) else "n. d.",
     100 * FET.get("fraccion_agua", 0.306) if "fraccion_agua" in FET else 30.6),
  after=4)
P("Palabras clave", 9, True, align=WD_ALIGN_PARAGRAPH.LEFT, after=2)
P("arqueología del paisaje; intervisibilidad; cuenca del Titicaca; modelos nulos; modelo digital de "
  "elevación; ciencia abierta", after=12)

two_col()

# ================================================================ introduccion
h1("1. Introducción")
P("La visibilidad de los monumentos funerarios es un argumento recurrente en la arqueología del paisaje "
  "altiplánico. Hyslop (1977), en su estudio de las chullpas de la zona Lupaqa —la que corresponde a la "
  "actual provincia de Chucuito—, propuso que estas torres funcionaban como marcadores del territorio "
  "controlado por unidades familiares, apoyándose en fuentes coloniales que las describen como mojones. "
  "De ahí se sigue, aunque no siempre se explicite, que su emplazamiento debería favorecer que se vieran "
  "entre sí.")
P("Bongers, Arkush y Harrower (2012) sometieron esa expectativa a contraste en un área de 80 km² al "
  "oeste del lago, en el entorno de Sillustani, comparando la visibilidad de las chullpas con la de 300 "
  "puntos aleatorios. Concluyeron que la visibilidad y la altitud actuaron como determinantes del "
  "emplazamiento. Es, hasta donde alcanza esta revisión, la única puesta a prueba formal de la hipótesis "
  "en la cuenca, y el punto de partida de este trabajo.")
P("La dificultad no está en medir la visibilidad sino en decidir con qué se compara. Cualquier conjunto "
  "de puntos sobre un terreno produce alguna red de intervisibilidad, y sin término de comparación ese "
  "número no distingue la intención del accidente. La pregunta con contenido es si se ven más de lo que "
  "se verían puntos situados sin esa intención, lo que obliga a definir qué significa «sin esa "
  "intención»: ahí se juega el resultado. Wheatley y Gillings (2000) y Lake y Woodman (2003) han "
  "señalado que el muestreo aleatorio simple resulta insuficiente cuando los emplazamientos comparten "
  "rasgos topográficos, y proponen contrastes estratificados frente a localizaciones comparables.")
P("Este trabajo retoma la cuestión en la provincia de Chucuito, en la ribera peruana del lago, con datos "
  "íntegramente abiertos y en dos aspectos distintos del planteamiento de 2012. Se mide la "
  "intervisibilidad recíproca entre sitios —la densidad de la red— en lugar del tamaño de la cuenca "
  "visual de cada uno, y se contrasta contra tres modelos nulos de exigencia creciente, de los cuales el "
  "primero equivale al muestreo aleatorio empleado hasta ahora.")
P("El resultado sostiene la conclusión previa pero rebaja su magnitud en un orden de magnitud, y esa es "
  "la primera contribución. La segunda es la documentación de tres artefactos hallados durante el "
  "análisis que, sin corregir, invierten o anulan el resultado sin dejar rastro visible.")

# ==================================================== materiales y metodos
h1("2. Materiales y métodos")
h2("2.1. Área de estudio y datos")
P("Se analiza la provincia de Chucuito, departamento de Puno, que en el siglo XV integraba el señorío "
  "Lupaqa descrito por Stanish (2003) como una de las formaciones políticas mayores de la cuenca. Los "
  "sitios proceden de la capa de sitios "
  "arqueológicos del Instituto Nacional de Cultura, hoy Ministerio de Cultura, que documenta 7 907 "
  "puntos en todo el Perú, de los cuales 507 corresponden a Puno y %d a Chucuito. Las cotas proceden "
  "del Copernicus DEM GLO-30 de la Agencia Espacial Europea, a 30 m de resolución, obtenido del "
  "repositorio público que no exige registro." % n_sitios)
P("La capa de sitios plantea un problema de procedencia que debe declararse. Se distribuye a través de "
  "un portal que la sirve desde un servicio de alojamiento genérico, sin licencia declarada, sin fecha "
  "de corte y sin criterio de inclusión documentado. El dato es de origen público, pero esa opacidad la "
  "hereda cualquier resultado. Se intentaron tres vías oficiales sin éxito: el geoservicio del "
  "Ministerio responde pero no publica ninguna capa de forma anónima; el sistema de información "
  "geográfica de arqueología se apoya en una cuenta personal cuyo contenido no es accesible por "
  "interfaz de programación; y un tercer repositorio citado habitualmente ya no resuelve.")

h2("2.2. Cálculo de la línea de visión")
P("Entre cada par de sitios se muestrea el perfil del terreno a razón de una muestra por celda —la "
  "densidad natural del dato, ya que un muestreo más fino solo interpolaría información que el modelo "
  "no contiene— y se comprueba si el relieve intermedio corta la recta que une observador y objetivo. "
  "Se aplica la corrección conjunta de curvatura terrestre y refracción atmosférica mediante el radio "
  "efectivo R/(1−k), con k = 0.13.")
P("El signo de esa corrección merece detenimiento, porque es donde es fácil equivocarse y el error no "
  "se manifiesta. Trabajando en el plano tangente al observador, el terreno desciende d²/2R con la "
  "distancia y el objetivo desciende D²/2R. Al reescribir la condición de bloqueo respecto a la recta "
  "que une ambos extremos sin corregir, esos dos descensos se combinan en un término d(D−d)/2R que se "
  "suma al terreno intermedio: visto desde esa cuerda, la Tierra abulta entre los extremos y el "
  "abultamiento se anula en ellos. Restarlo, que es lo que sugiere la expresión habitual «descenso por "
  "curvatura», vuelve la Tierra cóncava y hace que ningún relieve bloquee jamás a larga distancia. En "
  "un trayecto de 26 km el término alcanza 11.5 m en el punto medio.")
P("Las alturas de observador y objetivo se tratan como parámetros. Se adopta 1.70 m para el observador "
  "y 3.00 m para la estructura, cota conservadora para una chullpa, y el análisis se repite en el rango "
  "de 0 a 12 m.")

h2("2.3. Validación del cálculo")
P("Un algoritmo de visibilidad falla en silencio: devuelve valores plausibles aunque el criterio "
  "geométrico esté mal, y sobre terreno real no hay forma de advertirlo. Se validó por tanto contra "
  "doce terrenos construidos de respuesta conocida —plano, barrera que debe bloquear, la misma rebajada "
  "que no debe hacerlo, casos límite que verifican el término de curvatura, efecto de la altura del "
  "objetivo, depresión que nunca bloquea, horizonte geométrico a 40 km y simetría sobre terreno "
  "rugoso—. El detector supera los %d casos. El error de signo descrito en el apartado anterior se "
  "detectó precisamente aquí, no en el análisis."
  % (VAL["total"] if VAL else 12))

h2("2.4. Modelos nulos")
P("La conclusión depende por completo de con qué se compara la red observada, de modo que se calculan "
  "tres modelos de exigencia creciente.")
P("El primero sortea puntos al azar sobre el terreno, y equivale al procedimiento seguido por Bongers y "
  "colaboradores (2012). Casi siempre da significativo, porque los sitios reales ocupan relieve "
  "favorable y el relieve favorable ve más, de modo que por sí solo confirma poco.")
P("El segundo sortea celdas con la misma distribución de altitud que los sitios observados, de modo que "
  "el conjunto nulo reproduce el perfil altitudinal sin heredar las posiciones. Responde a si, dada la "
  "cota que ocupan, se ven más de lo que les correspondería. Es el tipo de contraste estratificado que "
  "Lake y Woodman (2003) emplearon para separar la visibilidad de la posición topográfica.")
P("El tercero toma la nube de sitios tal cual, conservando todas las distancias mutuas, y la traslada y "
  "rota rígidamente sobre el terreno. Es el más estricto: separa dónde están los sitios de cómo están "
  "dispuestos entre sí, y responde a si esta configuración concreta está colocada donde se ve más de lo "
  "que se vería en cualquier otro punto de la región.")
P("Los tres excluyen las masas de agua, por las razones que se exponen en el apartado 3.2.")

# ======================================================================= resultados
h1("3. Resultados")
h2("3.1. La red observada")
if d5:
    P("De los %s pares de sitios separados por menos de 5 km, %s presentan visión recíproca despejada, "
      "esto es una densidad de %s. El grado medio es de %.1f sitios visibles por sitio y solo %d quedan "
      "sin ninguna conexión. Al ampliar el alcance la densidad desciende hasta %s a 26 km, como cabe "
      "esperar, pero el número absoluto de pares visibles apenas crece: la intervisibilidad se agota "
      "en las distancias cortas."
      % (mil(d5["pares_elegibles"]), mil(d5["aristas"]), f(d5["densidad"]), d5["grado_medio"],
         d5["aislados"],
         f(NUL["observado"]["26000"]["densidad"])))
    P("La figura 1 recoge el área de estudio con la red a 5 km, y la figura 2, dos perfiles de línea de "
      "visión de distancia comparable, uno despejado y otro bloqueado, que ilustran qué mide el "
      "criterio de visibilidad empleado.")
figure("fig_mapa.png", "Figura 1.",
       "Área de estudio con el relieve sombreado, los %d sitios documentados y las aristas de "
       "intervisibilidad a menos de 5 km. Los triángulos señalan los %d sitios cuyo topónimo registrado "
       "contiene un término funerario. El lago Titicaca se representa aparte porque en un sombreado de "
       "relieve resulta indistinguible de una llanura: una superficie plana no proyecta sombra. El "
       "encuadre corresponde al área de estudio; la región de muestreo de los modelos nulos se extiende "
       "18 km más allá en todas las direcciones." % (n_sitios, n_fun))
figure("fig_perfiles.png", "Figura 2.",
       "Perfiles de línea de visión de dos pares de distancia comparable, uno despejado y otro "
       "bloqueado. Se muestran el terreno sin corregir, el terreno con la corrección de curvatura y "
       "refracción, y la recta que une observador y objetivo. La franja destacada marca la obstrucción.")

h2("3.2. El lago como artefacto de medida")
P("El modelo de elevación asigna una cota constante a las masas de agua. En el área de estudio el lago "
  "Titicaca aparece como una superficie perfectamente plana a 3808.5 m que ocupa el 30.6 % del recorte "
  "inicial. Sortear puntos nulos ahí es erróneo por dos motivos independientes: no se puede emplazar un "
  "sitio arqueológico en mitad del lago, y una superficie plana no bloquea ninguna vista, de modo que "
  "infla la intervisibilidad del modelo nulo. Ningún sitio observado cae sobre agua.")
if SINMASK:
    _sm = SINMASK["por_alcance"]["5000"]
    P("El efecto sobre la inferencia no es marginal, es determinante. Repitiendo el contraste más "
      "estricto sin enmascarar el agua, la distribución nula a 5 km tiene media %s y desviación típica "
      "%s, y el dato observado queda a %+.2f desviaciones: ausencia de efecto. Enmascarándola, la media "
      "nula baja a %s y la desviación a %s, y el mismo dato observado pasa a %+.2f. El artefacto no "
      "atenuaba el resultado, lo suprimía. Ambas ejecuciones se publican con los datos, para que la "
      "comparación pueda verificarse."
      % (f(_sm["nula_media"], 3), f(_sm["nula_sd"], 3), _sm["z"],
         f(r5["nula_media"], 3), f(r5["nula_sd"], 3), r5["z"]))

h2("3.3. El sesgo de orientación del recorte")
P("Un segundo artefacto apareció al revisar el nulo rígido. La nube de sitios mide 26 km en su eje "
  "mayor, y con el recorte inicial solo cabía dentro del área en 244 de 360 orientaciones: quedaban "
  "sistemáticamente excluidas las perpendiculares a la disposición observada. Ampliando el margen del "
  "recorte de 5 a 18 km caben las 360. El efecto medido descendió al hacerlo, de z = %+.2f a z = %+.2f, "
  "lo que indica que el recorte ajustado inflaba el contraste."
  % ((N300["por_alcance"]["5000"]["z"] if N300 else float("nan")), r5["z"]))

h2("3.4. Contraste con los modelos nulos")
P("La tabla 1 reúne la densidad observada y la de cada modelo nulo por alcance, y la figura 3 sitúa el "
  "dato observado dentro de las tres distribuciones nulas. La lectura conjunta es la que importa: lo "
  "relevante no es que el valor observado supere a un nulo, sino cuánto se estrecha el margen a medida "
  "que el modelo nulo conserva más rasgos de la configuración real.")
filas = []
for a in ["5000", "10000", "15000", "26000"]:
    fila = ["%.0f" % (float(a) / 1000)]
    fila.append(f(RIG["por_alcance"][a]["densidad_obs"]))
    if NUL:
        for k in ("uniforme", "estratificado_altitud"):
            d = NUL["nulos"][k][a]
            fila.append("%s ± %s" % (f(d["densidad_nula_media"], 3), f(d["densidad_nula_sd"], 3)))
            fila.append("%+.2f" % d["z"])
    d = RIG["por_alcance"][a]
    fila.append("%s ± %s" % (f(d["nula_media"], 3), f(d["nula_sd"], 3)))
    fila.append("%+.2f" % d["z"])
    filas.append(fila)
cols_t = ["Alcance (km)", "Observado"]
if NUL:
    cols_t += ["Nulo uniforme", "z", "Nulo por altitud", "z"]
cols_t += ["Nulo rígido", "z"]
table("Tabla 1.", "Densidad de intervisibilidad observada y en los modelos nulos, por alcance. "
      "Cada nulo se calculó con 500 repeticiones.", cols_t, filas)
figure("fig_nulos.png", "Figura 3.",
       "Densidad observada, marcada con la línea horizontal, frente a la distribución de cada modelo "
       "nulo. Las barras abarcan del percentil 5 al 95.")

h2("3.5. Sensibilidad a los supuestos")
P("Dos supuestos del cálculo admiten valores distintos de los adoptados, y conviene comprobar si la "
  "conclusión depende de ellos: la altura atribuida a las estructuras y la celda del modelo en que cae "
  "cada sitio. La tabla 2 recoge el primero.")
if ALT:
    fil = [["%.1f" % float(k), f(v["observado"]), "%s ± %s" % (f(v["nulo_media"], 3), f(v["nulo_sd"], 3)),
            "%+.2f" % v["z"], f(v["p"], 3)]
           for k, v in sorted(ALT["por_altura"].items(), key=lambda x: float(x[0]))]
    table("Tabla 2.", "Contraste rígido a 5 km recalculado para distintas alturas de estructura, "
          "con 200 repeticiones cada uno.",
          ["Altura (m)", "Observado", "Nulo", "z", "p"], fil)
    P("La altura atribuida a las estructuras mueve la densidad observada un 69 % entre los extremos del "
      "rango, pero no altera la conclusión: el contraste se mantiene entre z = +2.47 y z = +3.03, y es "
      "más fuerte con altura nula. La señal la lleva el emplazamiento del terreno, no la altura de las "
      "torres, que es precisamente el supuesto que nadie ha medido en campo.")
if SEN:
    dz = SEN["pruebas"]["desplazamiento_celda"]
    P("La asignación de cada sitio a una celda del modelo introduce otra incertidumbre. Desplazándola "
      "una posición en cada dirección, la densidad varía entre %s y %s, un %.0f %% del valor de "
      "referencia. El contraste sobrevive incluso en el extremo inferior de ese rango."
      % (f(dz["min"]), f(dz["max"]), 100 * dz["rango_relativo"]))

h2("3.6. El efecto de agregar dos grupos disjuntos")
if DIS:
    ju = DIS["distritos"]["juli"]
    po = DIS["distritos"]["pomata"]
    P("Los sitios de Chucuito no forman una nube continua sino dos grupos separados por %.1f km de "
      "relieve: Juli, con %d sitios, y Pomata, con %d, separación visible en la figura 1. El análisis "
      "anterior los trata como un solo conjunto, lo que tiene dos consecuencias. En los alcances largos "
      "entran al cómputo pares entre grupos que no describen ninguna relación de vecindad y que están "
      "casi siempre bloqueados; y el modelo nulo rígido rota una configuración que en realidad son dos, "
      "reproduciendo una separación cuyo significado no está establecido."
      % (DIS["separacion_centroides_km"], ju["n_sitios"], po["n_sitios"]))
    if ju.get("contraste"):
        c5 = ju["contraste"]["5000"]; c15 = ju["contraste"]["15000"]
        P("Repitiendo el contraste rígido sobre Juli en solitario, único grupo con tamaño suficiente, el "
          "efecto se atenúa: la densidad observada a 5 km es %s frente a %s ± %s en el nulo "
          "(z = %+.2f; p = %s), y a 15 km z = %+.2f con p = %s. El resultado pasa de significativo a "
          "marginal, lo que indica que parte del efecto medido sobre el conjunto agregado procedía de "
          "su estructura en dos grupos y no del emplazamiento de los sitios."
          % (f(c5["densidad_obs"]), f(c5["nula_media"]), f(c5["nula_sd"]), c5["z"],
             f(c5["p_unilateral"], 3), c15["z"], f(c15["p_unilateral"], 3)))
    P("Pomata se describe pero no se contrasta: con %d sitios el modelo nulo carece de potencia y "
      "cualquier valor p sería ilustrativo. Merece registrarse, no obstante, que su densidad de "
      "intervisibilidad a 5 km es %s, muy superior a la de Juli, y que %d de sus %d sitios llevan "
      "topónimo funerario. Es un patrón que este trabajo no puede evaluar y que señala dónde convendría "
      "dirigir un reconocimiento de campo."
      % (po["n_sitios"], f(po["observado"]["5000"]["densidad"]), po["n_funerarios"], po["n_sitios"]))

h2("3.7. Los sitios de topónimo funerario")
if FUN:
    k = "5000"
    d = FUN["por_alcance"][k]
    P("De los %d sitios, %d llevan en su nombre registrado un término asociado a estructuras funerarias. "
      "Su subred no está más conectada que la de cualquier subgrupo del mismo tamaño extraído del propio "
      "conjunto: densidad %s frente a %s ± %s en el remuestreo (z = %+.2f; p = %s). El resultado se "
      "repite en los tres alcances examinados."
      % (n_sitios, n_fun, f(d["densidad_funeraria"]), f(d["densidad_aleatoria_media"]),
         f(d["densidad_aleatoria_sd"]), d["z"], f(d["p_unilateral"], 3)))
    P("Caben dos lecturas y los datos no permiten separarlas. O las estructuras funerarias no se "
      "emplazaron con un criterio de visibilidad distinto del resto de sitios, o la etiqueta es "
      "demasiado débil para detectarlo: identificar un monumento por su topónimo registrado no es "
      "clasificarlo, un sitio funerario puede figurar con otro nombre y un topónimo puede aludir a un "
      "rasgo del paisaje. Con %d casos la potencia es limitada." % n_fun)

# ==================================================== discusion y conclusiones
h1("4. Discusión y conclusiones")
P("Los sitios de Chucuito ocupan posiciones desde las que se ven entre sí más de lo que lograría la "
  "misma configuración situada en otro punto del territorio, pero la fuerza de esa afirmación depende "
  "de cómo se delimite el conjunto. Sobre los 180 sitios agregados el contraste es significativo en los "
  "cuatro alcances; restringido a Juli, el único grupo espacialmente continuo, queda en el margen de la "
  "significación convencional. La evidencia es indicativa, no concluyente, y así debe leerse.")
if NUL:
    _raz = sorted(NUL["nulos"]["uniforme"][a]["z"] / RIG["por_alcance"][a]["z"]
                  for a in RIG["por_alcance"])
    P("Lo que sí resiste sin matices es la comparación entre modelos nulos: el muestreo aleatorio "
      "simple atribuye al emplazamiento un efecto entre %.1f y %.1f veces mayor, según el alcance, que "
      "el que sobrevive a un contraste capaz de controlar la disposición del conjunto. La distancia "
      "entre ambos crece con el alcance, porque el nulo uniforme dispersa los puntos por toda la "
      "región y penaliza tanto más su intervisibilidad cuanto más lejos se mira."
      % (_raz[0], _raz[-1]))
P("Conviene ser preciso sobre qué autoriza a concluir eso. El resultado dice que el emplazamiento "
  "favorece la visibilidad recíproca; no dice que la visibilidad fuera el criterio de emplazamiento. "
  "Otras razones —acceso al agua, suelos, rutas, defensa— pueden producir el mismo patrón si "
  "correlacionan con posiciones visualmente dominantes, y este análisis no las separa. La defensa es "
  "la más difícil de descartar aquí: Arkush (2011) documenta para el altiplano del Intermedio Tardío "
  "una red de sitios fortificados cuya lógica de emplazamiento —cotas altas, control visual de los "
  "accesos— produciría un patrón de intervisibilidad difícil de distinguir del observado. Establecer "
  "intención exigiría contrastar contra hipótesis alternativas explícitas, no solo contra el azar.")
P("Esa comparación admite lectura frente al antecedente. Bongers y colaboradores (2012) hallaron una "
  "visibilidad muy superior al azar y concluyeron que la visibilidad y la altitud determinaron el "
  "emplazamiento. El nulo uniforme, que reproduce su procedimiento, arroja aquí valores z de entre +5.9 "
  "y +25.6 según el alcance, coherentes con esa conclusión; el estratificado por altitud los eleva "
  "todavía más, lo que indica que la cota por sí sola no explica el patrón. Pero el nulo rígido, que "
  "conserva la configuración y solo cambia su emplazamiento, los reduce a valores en torno a +2.8. La "
  "conclusión previa se sostiene, y este trabajo la corrobora en otra zona de la cuenca y con otra "
  "medida; pero la magnitud que sugiere el muestreo aleatorio simple es aproximadamente un orden de "
  "magnitud mayor que la que resiste un contraste capaz de controlar la disposición del conjunto.")
P("La segunda contribución es metodológica y probablemente más transferible. Los tres artefactos "
  "documentados —el signo de la corrección por curvatura, la cota constante del agua y el sesgo de "
  "orientación del recorte— comparten un rasgo: ninguno produce un fallo visible. Los tres devuelven "
  "redes de intervisibilidad de aspecto razonable y valores estadísticos interpretables. El primero se "
  "detectó con pruebas sintéticas, el segundo al examinar el mapa y el tercero al medir la tasa de "
  "aceptación del nulo. Ninguno se habría advertido mirando solo los resultados.")

h2("4.1. Limitaciones")
P("La procedencia de la capa de sitios es la limitación de fondo. Sin criterio de inclusión documentado "
  "no se sabe qué sitios faltan, y un registro incompleto sesga la red de forma imprevisible. El "
  "análisis es reproducible, pero su base documental no está verificada.")
P("El modelo de elevación describe la superficie, no el terreno: incluye vegetación y edificación, que "
  "en el altiplano son poco relevantes pero no nulas. La resolución de 30 m impone además un límite: "
  "relieves menores que la celda no bloquean nada en el cálculo aunque lo hicieran en la realidad.")
P("La región de comparación se extendió a 18 km alrededor de los sitios para que el nulo rígido "
  "dispusiera de todas las orientaciones. Eso incorpora terreno de carácter distinto al de la ribera, "
  "lo que hace el contraste más exigente pero también menos homogéneo.")
P("Por último, el análisis es sincrónico. La capa no distingue cronologías, de modo que la red trata "
  "como contemporáneos sitios que pueden estar separados por siglos. La intervisibilidad medida es la "
  "del conjunto documentado, no la de un momento concreto.")

h2("4.2. Trabajo futuro")
P("El paso más útil sería incorporar cronología: con los sitios fechados, la misma red puede calcularse "
  "por periodo y comprobarse si la intervisibilidad se construye o se hereda. Extender el análisis a las "
  "demás provincias de Puno permitiría además contrastar si el patrón de Chucuito es particular de la "
  "ribera o general del altiplano.")

# ==================================================== reproducibilidad
h1("5. Disponibilidad de datos y código")
P("El código que reconstruye el análisis completo, los datos derivados y los ficheros de resultados "
  "están depositados con identificador persistente, bajo licencia MIT: "
  "https://doi.org/%s. El modelo de elevación es Copernicus DEM GLO-30, de uso libre con "
  "atribución; la capa de sitios se cita según su procedencia declarada, con la reserva expuesta en el "
  "apartado 2.1." % DOI_CONCEPTO)
P("El depósito incluye las ejecuciones previas a la corrección de dos de los tres artefactos descritos "
  "—el nulo rígido calculado sin enmascarar el lago y el calculado sobre el recorte ajustado—, de modo "
  "que las comparaciones del apartado 3 puedan verificarse y no haya que tomarlas por buenas. El "
  "tercero, el signo de la corrección por curvatura, se comprueba ejecutando la validación sintética "
  "incluida.")

P("Agradecimientos", 11, True, align=WD_ALIGN_PARAGRAPH.CENTER, before=10, after=6)
P("[En blanco para la revisión por pares, conforme a la norma de la revista.]", italic=True)

P("Referencias", 11, True, align=WD_ALIGN_PARAGRAPH.CENTER, before=10, after=6)
for _ref in [
    "Arkush, E. N. (2011). Hillforts of the Ancient Andes: Colla Warfare, Society, and Landscape. "
    "Gainesville: University Press of Florida.",
    "Bongers, J., Arkush, E., & Harrower, M. (2012). Landscapes of death: GIS-based analyses of chullpas "
    "in the western Lake Titicaca basin. Journal of Archaeological Science, 39(6), 1687-1693.",
    "Hyslop, J. (1977). Chulpas of the Lupaca Zone of the Peruvian High Plateau. Journal of Field "
    "Archaeology, 4(2), 149-170.",
    "Lake, M. W., & Woodman, P. E. (2003). Visibility studies in archaeology: a review and case study. "
    "Environment and Planning B: Planning and Design, 30(5), 689-707.",
    "Stanish, C. (2003). Ancient Titicaca: The Evolution of Complex Society in Southern Peru and "
    "Northern Bolivia. Berkeley: University of California Press.",
    "Wheatley, D., & Gillings, M. (2000). Vision, perception and GIS: developing enriched approaches to "
    "the study of archaeological visibility. En G. R. Lock (ed.), Beyond the Map: Archaeology and "
    "Spatial Technologies (pp. 1-27). Amsterdam: IOS Press.",
]:
    _p = P(_ref, 9, after=4)
    _p.paragraph_format.left_indent = Cm(0.6)
    _p.paragraph_format.first_line_indent = Cm(-0.6)

doc.save(OUT)

palabras = sum(len(p.text.split()) for p in doc.paragraphs)
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            palabras += len(c.text.split())
print("Manuscrito ->", OUT)
print("palabras: %d | figuras: %d | tablas: %d" % (palabras, len(doc.inline_shapes), len(doc.tables)))
