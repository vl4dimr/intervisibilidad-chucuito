# -*- coding: utf-8 -*-
"""
Auditoria del manuscrito: formato, figuras y coherencia numerica.

El manuscrito se genera leyendo los JSON de resultados, de modo que las cifras
insertadas por codigo no pueden desviarse. Pero hay valores escritos a mano en
el texto —los que redondean, comparan versiones o resumen una tendencia— y esos
si envejecen cuando algo se recalcula. Son la fuente habitual de que un articulo
acabe afirmando un numero que ya no es suyo.

Esta auditoria comprueba tres cosas:

  formato    tipografia, cuerpos, columnas, anonimato y marcador decimal, contra
             la especificacion de Virtual Archaeology Review
  figuras    existencia, resolucion, y que cada una este citada en el texto
  contenido  cada cifra del texto que deberia proceder de un resultado, cotejada
             contra el JSON correspondiente
  prosa      referencias listadas pero no citadas, tildes perdidas al concatenar
             cadenas, tablas sin anunciar y un resumen que afirme mas de lo que
             sostienen los resultados
  plantilla  cotejo contra VAR_Template.dot version 13, descargada del portal de
             la revista: margenes, resumenes en ambos idiomas, extended abstract
             obligatorio, limites de highlights y palabras clave, cuerpos y DOI

Salida: results/auditoria.json
"""
import json
import os
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
RES = os.path.join(BASE, "results")
FIG = os.path.join(RES, "figuras")
DOCX = os.path.join(BASE, "VAR_intervisibilidad.docx")

ok_n, fail_n = 0, 0
items = []


def check(bloque, etiqueta, cond, detalle=""):
    global ok_n, fail_n
    if cond:
        ok_n += 1
    else:
        fail_n += 1
    items.append({"bloque": bloque, "check": etiqueta, "ok": bool(cond), "detalle": detalle})
    print("  %-5s %-54s %s" % ("OK" if cond else "FALLA", etiqueta, detalle), flush=True)


def cargar(n):
    p = os.path.join(RES, n)
    return json.load(open(p, encoding="utf-8")) if os.path.exists(p) else None


def main():
    if not os.path.exists(DOCX):
        sys.exit("Falta el manuscrito")
    doc = Document(DOCX)
    ps = [p.text for p in doc.paragraphs]
    texto = "\n".join(ps)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                texto += "\n" + c.text

    NUL, RIG = cargar("nulos.json"), cargar("nulo_rigido.json")
    FUN, VAL = cargar("funerarios.json"), cargar("validacion_los.json")
    DIS, ALT = cargar("por_distrito.json"), cargar("nulo_alturas.json")
    SEN = cargar("sensibilidad.json")
    TER = json.load(open(os.path.join(BASE, "data", "terreno_log.json"), encoding="utf-8"))

    print("AUDITORÍA DEL MANUSCRITO\n")

    # ---------------------------------------------------------------- formato
    print("1. FORMATO")
    runs = [r for p in doc.paragraphs for r in p.runs if r.text.strip()]
    fuentes = {r.font.name for r in runs}
    heredadas = sum(1 for r in runs if r.font.name is None)
    check("formato", "fuente única Arial", fuentes <= {"Arial"} and heredadas == 0,
          "fuentes %s, heredadas %d" % (sorted(x for x in fuentes if x), heredadas))

    colores = set()
    for r in runs:
        try:
            if r.font.color is not None and r.font.color.rgb is not None:
                colores.add(str(r.font.color.rgb))
        except Exception:
            pass
    check("formato", "sin color distinto del negro", colores <= {"000000"},
          "colores: %s" % (sorted(colores) or ["(por defecto)"]))

    cols = []
    for sec in doc.sections:
        c = sec._sectPr.xpath("./w:cols")[0]
        cols.append(int(c.get(qn("w:num")) or 1))
        w = (sec.page_width - sec.left_margin - sec.right_margin) / 36000.0
    check("formato", "caja de texto de 170 mm", abs(w - 170) < 1.5, "%.0f mm" % w)
    check("formato", "cuerpo a dos columnas", 2 in cols, "secciones: %s" % len(cols))

    cuerpo = [p for p in doc.paragraphs
              if len(p.text.split()) > 25
              and p.style.name != "List Bullet"   # las viñetas no son texto corrido
              and not re.match(r"^(Figura|Tabla)\s+\d+\.", p.text.strip())
              and not p.text.strip().startswith("[Fichero anónimo")]
    tam = {r.font.size.pt for p in cuerpo for r in p.runs if r.font.size}
    check("formato", "texto corrido a cuerpo 9", tam == {9.0}, "cuerpos: %s" % sorted(tam))
    just = sum(1 for p in cuerpo if p.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY)
    check("formato", "texto justificado", just == len(cuerpo), "%d de %d" % (just, len(cuerpo)))

    comas = re.findall(r"\d,\d", texto)
    check("formato", "marcador decimal con punto", not comas,
          "" if not comas else "%d comas decimales: %s" % (len(comas), sorted(set(comas))[:6]))

    nombres = ["Mamani", "Alanoca", "Duany", "Vicente A", "Altiplano"]
    hay = [x for x in nombres if x in texto]
    check("formato", "fichero anónimo", not hay, "aparece: %s" % hay if hay else "")

    # ---------------------------------------------------------------- figuras
    print("\n2. FIGURAS")
    from PIL import Image
    esperadas = ["fig_mapa.png", "fig_perfiles.png", "fig_nulos.png"]
    for fn in esperadas:
        p = os.path.join(FIG, fn)
        if not os.path.exists(p):
            check("figuras", fn, False, "no existe")
            continue
        im = Image.open(p)
        # PNG almacena pixeles por metro en entero: 300 dpi se relee como 299.
        dpi = im.info.get("dpi", (0, 0))[0]
        check("figuras", "%s: resolución" % fn, dpi >= 299 and im.width >= 1800,
              "%dx%d px, %d dpi" % (im.width, im.height, dpi))
    check("figuras", "tres figuras incrustadas", len(doc.inline_shapes) == 3,
          "%d incrustadas" % len(doc.inline_shapes))
    cuerpo_txt = "\n".join(p for p in ps if not p.strip().startswith("Figura "))
    for i in (1, 2, 3):
        check("figuras", "Figura %d citada en el texto" % i, ("figura %d" % i) in cuerpo_txt.lower())
    check("figuras", "pies numerados y presentes",
          len([p for p in ps if re.match(r"^Figura\s+\d+\.", p.strip())]) == 3)
    check("figuras", "tablas numeradas y presentes",
          len([p for p in ps if re.match(r"^Tabla\s+\d+\.", p.strip())]) == len(doc.tables),
          "%d pies, %d tablas" % (len([p for p in ps if re.match(r"^Tabla\s+\d+\.", p.strip())]),
                                  len(doc.tables)))

    # --------------------------------------------------------------- contenido
    print("\n3. COHERENCIA NUMÉRICA")

    def en_texto(valor, dec=2):
        """¿Aparece el valor en el texto, con el formato habitual del documento?"""
        formas = {("%%.%df" % d) % valor for d in range(dec, dec + 2)}
        formas |= {("%+.2f" % valor)}
        return any(x in texto for x in formas)

    check("contenido", "número de sitios", str(TER["sitios"]) in texto, "%d" % TER["sitios"])
    check("contenido", "sitios con topónimo funerario",
          str(TER["funerarios_por_toponimo"]) in texto, "%d" % TER["funerarios_por_toponimo"])
    if VAL:
        check("contenido", "casos de validación superados",
              ("%d" % VAL["total"]) in texto, "%d de %d" % (VAL["superados"], VAL["total"]))
    if RIG:
        r5 = RIG["por_alcance"]["5000"]
        check("contenido", "z del nulo rígido a 5 km", en_texto(r5["z"]), "%+.2f" % r5["z"])
        check("contenido", "densidad observada a 5 km", en_texto(r5["densidad_obs"], 4),
              "%.4f" % r5["densidad_obs"])
    if NUL:
        zu = [NUL["nulos"]["uniforme"][k]["z"] for k in ("5000", "26000")]
        check("contenido", "rango z del nulo uniforme citado",
              all(("%.1f" % z) in texto or ("%.2f" % z) in texto for z in zu),
              "de %+.2f a %+.2f" % (zu[0], zu[1]))
    if DIS and DIS["distritos"]["juli"].get("contraste"):
        j5 = DIS["distritos"]["juli"]["contraste"]["5000"]
        check("contenido", "z de Juli a 5 km", en_texto(j5["z"]), "%+.2f" % j5["z"])
        check("contenido", "p de Juli a 5 km", ("%.3f" % j5["p_unilateral"]) in texto,
              "%.3f" % j5["p_unilateral"])
        check("contenido", "separación entre centroides",
              ("%.1f" % DIS["separacion_centroides_km"]) in texto,
              "%.1f km" % DIS["separacion_centroides_km"])
    if FUN:
        f5 = FUN["por_alcance"]["5000"]
        check("contenido", "p de la subred funeraria", ("%.3f" % f5["p_unilateral"]) in texto,
              "%.3f" % f5["p_unilateral"])

    # cifras escritas a mano que suelen envejecer
    print("\n4. CIFRAS ESCRITAS A MANO")
    manuales = re.findall(r"z = ([+-]\d+\.\d+)", texto)
    print("     valores z citados en el texto: %s" % sorted(set(manuales)))
    vivos = set()
    if RIG:
        vivos |= {"%+.2f" % v["z"] for v in RIG["por_alcance"].values()}
    if NUL:
        for b in NUL["nulos"].values():
            vivos |= {"%+.2f" % v["z"] for v in b.values()}
    if DIS and DIS["distritos"]["juli"].get("contraste"):
        vivos |= {"%+.2f" % v["z"] for v in DIS["distritos"]["juli"]["contraste"].values()}
    if ALT:
        vivos |= {"%+.2f" % v["z"] for v in ALT["por_altura"].values()}
    if FUN:
        vivos |= {"%+.2f" % v["z"] for v in FUN["por_alcance"].values()}
    # Los dos contraejemplos: el nulo sin mascara de agua y el del recorte previo.
    # Sus cifras aparecen en el texto como termino de comparacion y necesitan
    # respaldo igual que cualquier otra.
    for extra in ("nulo_rigido_sin_mascara.json", "nulo_rigido_n300.json"):
        j = cargar(extra)
        if j:
            vivos |= {"%+.2f" % v["z"] for v in j["por_alcance"].values()}
    huerfanos = [m for m in set(manuales) if m not in vivos]
    check("contenido", "todo valor z citado existe en los resultados", not huerfanos,
          "sin respaldo: %s" % huerfanos if huerfanos else "")

    # ------------------------------------------------------- prosa y aparato
    # Estas comprobaciones nacen de fallos reales que la version anterior de la
    # auditoria no vio y que hubo que encontrar leyendo: referencias listadas
    # pero nunca citadas, palabras que perdieron la tilde al construirse la
    # cadena por concatenacion, y tablas que aparecian sin que el texto las
    # anunciara.
    print("\n5. PROSA Y APARATO CRÍTICO")

    partes = texto.split("Referencias")
    cuerpo_ref, lista_ref = (partes[0], partes[-1]) if len(partes) > 1 else (texto, "")
    citadas, sin_citar = [], []
    for linea in lista_ref.split("\n"):
        m = re.match(r"^([A-ZÁÉÍÓÚÑ][\w'áéíóúñ-]+),", linea.strip())
        if not m:
            continue
        (citadas if m.group(1) in cuerpo_ref else sin_citar).append(m.group(1))
    check("prosa", "toda referencia listada está citada", not sin_citar,
          "sin citar: %s" % sin_citar if sin_citar else "%d referencias" % len(citadas))

    # Palabras que el documento usa y que pierden la tilde con facilidad al
    # concatenar cadenas en el generador.
    ACENTOS = ["analisis", "configuracion", "separacion", "relacion", "estan",
               "rigido", "computo", "seccion", "region", "linea", "area",
               "numero", "metodo", "ademas", "asi", "aqui", "segun", "mas alla",
               "esta establecido", "tambien", "solo se", "deposito"]
    # Solo sobre la parte en espanol: el resumen en ingles, el extended abstract
    # y las referencias usan legitimamente «area», «region» o «section».
    es = texto
    for ini, fin in (("Abstract", "Resumen"), ("Extended abstract", "\n1. "), ("Referencias", None)):
        i = es.find(ini)
        if i < 0:
            continue
        j = es.find(fin, i) if fin else len(es)
        es = es[:i] + (es[j:] if j > 0 else "")
    perdidas = [w for w in ACENTOS if re.search(r"\b%s\b" % w, es)]
    check("prosa", "sin palabras que hayan perdido la tilde", not perdidas,
          "revisar: %s" % perdidas if perdidas else "")

    for i in range(1, len(doc.tables) + 1):
        cuerpo_tab = "\n".join(p for p in ps if not p.strip().startswith("Tabla "))
        check("prosa", "Tabla %d citada en el texto" % i,
              ("tabla %d" % i) in cuerpo_tab.lower())

    # Enteros de cuatro cifras o mas: o todos con separador de millar o ninguno.
    # Se excluyen anos y paginas de las referencias, que no lo llevan.
    sin_sep = set(re.findall(r"(?<![.\d\u202f])\d{4,}(?![.\d])", cuerpo_ref))
    sin_sep = {x for x in sin_sep if not (1500 < int(x) < 2100)}
    check("prosa", "separador de millar uniforme", not sin_sep,
          "sin separador: %s" % sorted(sin_sep) if sin_sep else "")

    # El resumen no debe afirmar mas de lo que sostiene el apartado 3.6.
    resumen = texto.split("Palabras clave")[0]
    if DIS and DIS["distritos"]["juli"].get("contraste"):
        pj = DIS["distritos"]["juli"]["contraste"]["5000"]["p_unilateral"]
        check("prosa", "el resumen recoge la atenuación por grupos",
              ("%.3f" % pj) in resumen, "p de Juli = %.3f" % pj)

    # ------------------------------------------------ plantilla oficial VAR
    # Reglas tomadas de VAR_Template.dot (version 13), descargada del portal de
    # la revista, no de una especificacion reconstruida de memoria. Las cuatro
    # que fallaban al cotejar por primera vez fueron el margen superior, el
    # resumen en ingles, el extended abstract y los DOI de las referencias.
    print("\n6. PLANTILLA OFICIAL DE VAR")

    sec0 = doc.sections[0]
    mm = lambda v: v / 36000.0
    margenes = (mm(sec0.top_margin), mm(sec0.bottom_margin),
                mm(sec0.left_margin), mm(sec0.right_margin))
    check("plantilla", "márgenes 25 / 20 / 20 / 20 mm",
          all(abs(a - b) < 1 for a, b in zip(margenes, (25, 20, 20, 20))),
          "sup %.0f inf %.0f izq %.0f der %.0f" % margenes)

    def bloque(inicio, fin):
        """Texto entre dos encabezados del preliminar."""
        try:
            i, j = ps.index(inicio), ps.index(fin)
        except ValueError:
            return None
        return " ".join(ps[i + 1:j])

    limpio = [p.strip() for p in ps]
    ps_ = limpio

    def entre(a, b):
        try:
            return " ".join(ps_[ps_.index(a) + 1:ps_.index(b)])
        except ValueError:
            return None

    ab = entre("Abstract", "Keywords")
    check("plantilla", "resumen en inglés, hasta 300 palabras",
          ab is not None and len(ab.split()) <= 300,
          "%d palabras" % len(ab.split()) if ab else "AUSENTE")

    res = entre("Resumen", "Palabras clave")
    check("plantilla", "resumen en español, hasta 300 palabras",
          res is not None and len(res.split()) <= 300,
          "%d palabras" % len(res.split()) if res else "AUSENTE")

    # Obligatorio cuando el articulo va en espanol.
    ext = None
    for cierre in ps_:
        if cierre.startswith("1. "):
            ext = entre("Extended abstract", cierre)
            break
    n_ext = len(ext.split()) if ext else 0
    check("plantilla", "extended abstract en inglés, de 600 a 900 palabras",
          600 <= n_ext <= 900, "%d palabras" % n_ext if ext else "AUSENTE")

    vinetas = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
    largos = [len(p.text) for p in vinetas]
    check("plantilla", "tres highlights de 190 caracteres como máximo",
          len(vinetas) == 3 and all(n <= 190 for n in largos),
          "%d viñetas, máximo %d caracteres" % (len(vinetas), max(largos) if largos else 0))

    for etiqueta, cab in (("inglés", "Keywords"), ("español", "Palabras clave")):
        try:
            linea = ps_[ps_.index(cab) + 1]
            n = len([x for x in linea.split(";") if x.strip()])
            check("plantilla", "hasta seis palabras clave en %s" % etiqueta, n <= 6, "%d" % n)
        except ValueError:
            check("plantilla", "hasta seis palabras clave en %s" % etiqueta, False, "AUSENTE")

    # Cuerpos: H1 a 11, H2 a 10, texto a 9, tablas y pies a 8.
    h1 = {p.runs[0].font.size.pt for p in doc.paragraphs
          if re.match(r"^\d+\.\s", p.text.strip()) and p.runs and p.runs[0].font.size}
    h2 = {p.runs[0].font.size.pt for p in doc.paragraphs
          if re.match(r"^\d+\.\d+\.\s", p.text.strip()) and p.runs and p.runs[0].font.size}
    check("plantilla", "encabezados de nivel 1 a cuerpo 11", h1 == {11.0}, "%s" % sorted(h1))
    check("plantilla", "encabezados de nivel 2 a cuerpo 10", h2 == {10.0}, "%s" % sorted(h2))

    t_tab = {r.font.size.pt for t in doc.tables for row in t.rows for c in row.cells
             for p in c.paragraphs for r in p.runs if r.font.size}
    check("plantilla", "tablas a cuerpo 8", t_tab == {8.0}, "%s" % sorted(t_tab))

    # APA 6.a con DOI cuando exista, en la forma «doi:».
    refs = [p for p in ps_ if re.match(r"^[A-Z][a-zA-Z\-]+,\s+[A-Z]\.", p)]
    con_doi = [r for r in refs if "doi:" in r]
    check("plantilla", "los DOI se escriben con el prefijo «doi:»",
          all(("http" not in r.split("doi:")[-1]) for r in con_doi),
          "%d de %d referencias con DOI" % (len(con_doi), len(refs)))

    # Marcadores de posicion: lo que se escribe «para rellenar luego» es
    # exactamente lo que acaba enviandose sin rellenar.
    corchetes = [p for p in ps if p.strip().startswith("[") and p.strip().endswith("]")]
    check("plantilla", "sin marcadores de posición entre corchetes", not corchetes,
          "quedan %d: %s" % (len(corchetes), corchetes[0][:50]) if corchetes else "")

    # El titulo en segunda lengua va alineado a la izquierda (estilo VAR Titulo).
    t2 = doc.paragraphs[1] if len(doc.paragraphs) > 1 else None
    check("plantilla", "título en segunda lengua a la izquierda y en cursiva",
          t2 is not None and t2.alignment == WD_ALIGN_PARAGRAPH.LEFT
          and t2.runs and t2.runs[0].italic,
          "%s" % (t2.text[:40] if t2 else "ausente"))

    # ------------------------------------------ propiedades del propio fichero
    # La revision es ciega, y las propiedades del .docx son el primer sitio
    # donde mira quien quiera saber quien firma. python-docx las rellena solo.
    print("\n7. PROPIEDADES DEL FICHERO")
    cp = doc.core_properties
    rastro = [k for k in ("author", "last_modified_by", "comments", "category",
                          "content_status", "identifier", "subject", "keywords")
              if (getattr(cp, k, "") or "").strip()]
    check("fichero", "sin autor ni rastro del generador", not rastro,
          "con contenido: %s" % rastro if rastro else "")

    import zipfile
    with zipfile.ZipFile(DOCX) as z:
        crudo = z.read("docProps/core.xml").decode("utf-8")
        crudo += z.read("docProps/app.xml").decode("utf-8")
    sospechas = [w for w in ("python-docx", "Macintosh") + tuple(nombres) if w in crudo]
    check("fichero", "los metadatos XML no delatan nada", not sospechas,
          "aparece: %s" % sospechas if sospechas else "")

    try:
        check("fichero", "fecha de creación verosímil", cp.created.year >= 2026,
              "%s" % cp.created.year)
    except Exception:
        check("fichero", "fecha de creación verosímil", False, "ilegible")

    print("\n" + "=" * 66)
    print("RESULTADO: %d comprobaciones correctas, %d fallos" % (ok_n, fail_n))
    print("=" * 66)
    json.dump({"ok": ok_n, "fallos": fail_n, "items": items},
              open(os.path.join(RES, "auditoria.json"), "w", encoding="utf-8"),
              ensure_ascii=False, indent=2)
    return 1 if fail_n else 0


if __name__ == "__main__":
    sys.exit(main())
